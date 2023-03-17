import pandas as pd
import numpy as np
from langid import classify
from sklearn.neural_network import MLPClassifier
import time
from laserembeddings import Laser
import warnings
import smbclient
import tempfile
from .excel_processing import *

# from environment import MODE
#
# if MODE == 'local':
#     from .local_constants import *
# else:
#     from .dev_constants import *

warnings.filterwarnings("ignore")
import joblib
from docx import Document
import re
from bs4 import BeautifulSoup
import mammoth

laser = Laser(path_to_bpe_codes, path_to_bpe_vocab, path_to_encoder)

nutrition_keywors = ['Energy', 'Calories', 'Protein', 'Carbohydrate', 'Sugars', 'Total Fat', 'Saturated Fat', 'Fibre',
                     'Sodium', 'Salt', 'Trans Fat']

nutri_keys_para = ['zinc', 'total fat', 'saturated fat', 'transfat', 'cholesterol', 'sodium', 'carbohydrate',
                   'dietary fiber', 'total sugars', 'includes added sugars', 'protein', 'servings per package',
                   'serving size', 'calories', 'servings per container', 'trans fat', 'vitamin d', 'vitamin c', 'zinc',
                   'salt']


classifier = joblib.load(mond_gen_model)
classifier_key = joblib.load(mond_nutri_model)

# document_location = r"/Users/sakthivelv/Documents/SGK/Mondelez_MEA/"
# temp_directory = tempfile.TemporaryDirectory(dir=document_location)
# input_docx_location = f'{temp_directory.name}/input_docx.docx'


def preprocess(file):
    t1 = time.time()
    #     document = Document(file)
    html = mammoth.convert_to_html(file).value
    soup = BeautifulSoup(html, "html.parser")
    table_content_list_all = []
    for tables in soup.find_all('table'):
        for row in tables.find_all('tr'):
            column_list = []
            for column in row.find_all('td'):
                #             column_list.append(str(column).replace('<td>','').replace('</td>','').replace('</p>','').replace('<p>','').replace('<td colspan="2">','').strip())
                raw_html = str(column).replace('<strong>', 'start_bold').replace('</strong>', 'end_bold').replace(
                    '</p>', '\n').strip()
                cleantext = BeautifulSoup(raw_html, "lxml").text
                column_list.append(
                    cleantext.replace('<', '&lt;').replace('>', '&gt;').replace('start_bold', '<b>').replace('end_bold',
                                                                                                             '</b>'))
            table_content_list_all.append(column_list)

    table_content_list = []
    for cnt in range(0, len(table_content_list_all)):
        sr = pd.Series(table_content_list_all[cnt])
        result = sr.first_valid_index()
        if sr[result].strip() != '':
            table_content_list.append(table_content_list_all[cnt])
        elif sr[result].strip() == '':
            content = str(('').join(table_content_list_all[cnt])).strip()
            if 'kcal' in content.lower():
                table_content_list_all[cnt][0] = 'Calories'
                table_content_list.append(table_content_list_all[cnt])

    thai_nutrition = []
    indo_nutrition = []
    en_nutrition = []
    final_content_list = []
    ingre_dic = {}
    nutri_para_list = []

    for k in range(0, len(table_content_list)):
        if "serving size" in table_content_list[k][0].lower() or "serving size declaration" in table_content_list[k][
            0].lower():
            if table_content_list[k][1:]:
                ingre_dic.setdefault("Serving Size", []).append(
                    {classify(table_content_list[k][1])[0]: table_content_list[k][1].strip()})
        classified_output = classifier.predict(laser.embed_sentences(
            table_content_list[k][0].replace('&lt;', '').replace('&gt;', '').replace('</b>', '').replace('<b>', ''),
            lang='en'))
        probability1 = classifier.predict_proba(laser.embed_sentences(
            table_content_list[k][0].replace('&lt;', '').replace('&gt;', '').replace('</b>', '').replace('<b>', ''),
            lang='en'))
        probability1.sort()
        prob1 = probability1[0][-1]
        if prob1 > 0.75:
            classified_output = classified_output[0]
        else:
            classified_output = 'UNMAPPED'

        if classified_output == 'THAI_NUTRITION':
            thai_nutrition.append(table_content_list[k])

        elif classified_output == 'INDO_NUTRITION':
            indo_nutrition.append(table_content_list[k])

        elif classified_output == 'EN_NUTRITION':
            if 'protein' in table_content_list[k][0].lower():
                classified_output = classifier.predict(laser.embed_sentences(
                    table_content_list[k - 1][0].replace('&lt;', '').replace('&gt;', '').replace('</b>', '').replace(
                        '<b>', ''), lang='en'))
                probability1 = classifier.predict_proba(laser.embed_sentences(
                    table_content_list[k - 1][0].replace('&lt;', '').replace('&gt;', '').replace('</b>', '').replace(
                        '<b>', ''), lang='en'))
                probability1.sort()
                prob1 = probability1[0][-1]
                if prob1 > 0.75:
                    classified_output = classified_output[0]
                else:
                    classified_output = 'UNMAPPED'
                if classified_output == 'INDO_NUTRITION':
                    indo_nutrition.append(table_content_list[k])
                elif classified_output == 'EN_NUTRITION':
                    en_nutrition.append(table_content_list[k])
            else:
                en_nutrition.append(table_content_list[k])

        elif classified_output not in ['None', 'INGREDIENTS_DECLARATION', 'NUTRITION_INSTRUCTIONS',
                                       'NUTRITION_PARAGRAPH', 'THAI_NUTRITION', 'INDO_NUTRITION', 'EN_NUTRITION',
                                       'ALLERGEN_STATEMENT']:
            contents = [table_content_list[k], classified_output.strip()]
            final_content_list.append(contents)

        elif classified_output in ['INGREDIENTS_DECLARATION', 'NUTRITION_INSTRUCTIONS']:
            lang = classify(table_content_list[k][0])[0]
            if classified_output in ingre_dic:
                ingre_dic[classified_output].append(
                    {lang: table_content_list[k][0].replace('<b> </b>', '').replace('\xa0', '').strip()})
            else:
                ingre_dic[classified_output] = [
                    {lang: table_content_list[k][0].replace('<b> </b>', '').replace('\xa0', '').strip()}]

        elif classified_output in ['NUTRITION_PARAGRAPH']:
            nutri_para_list.append(table_content_list[k][0])

        elif classified_output == 'ALLERGEN_STATEMENT':
            classified_output = 'INGREDIENTS_DECLARATION'
            lang = classify(table_content_list[k][0])[0]
            if classified_output in ingre_dic:
                ingre_dic[classified_output].append(
                    {lang: table_content_list[k][0].replace('<b> </b>', '').replace('\xa0', '')})
            else:
                ingre_dic[classified_output] = [
                    {lang: table_content_list[k][0].replace('<b> </b>', '').replace('\xa0', '')}]
        continue
    ingre_new_dic = {}
    for k1, v1 in ingre_dic.items():
        for cnts in v1:
            for lang, txt in cnts.items():
                split_txt = (txt.split('\n\n'))
                #             print('*'*10)
                for txts in split_txt:
                    if txts.strip() not in ['', '\n']:
                        classified_output = classifier.predict(
                            laser.embed_sentences(txts.replace('</b>', '').replace('<b>', ''), lang='en'))
                        probability1 = classifier.predict_proba(
                            laser.embed_sentences(txts.replace('</b>', '').replace('<b>', ''), lang='en'))
                        probability1.sort()
                        prob1 = probability1[0][-1]
                        if prob1 > 0.75:
                            classified_output = classified_output[0]
                        else:
                            classified_output = 'UNMAPPED'

                        if classified_output != 'None':
                            if classified_output in ingre_new_dic:
                                ingre_new_dic[classified_output].append({lang: txts})
                            else:
                                ingre_new_dic[classified_output] = [{lang: txts}]

                        elif classified_output == 'UNMAPPED':
                            if k1 in ingre_new_dic:
                                ingre_new_dic[k1].append({lang: txts})
                            else:
                                ingre_new_dic[k1] = [{lang: txts}]

    gen_cate_dic = {}
    for p in range(0, len(final_content_list)):
        if len(final_content_list[p][0]) > 1:
            for d in range(1, len(final_content_list[p][0])):
                #         for key in gen_list:
                if final_content_list[p][0][d].strip() not in ['Servings per Package (EACH):',
                                                               'Servings per Package (INNER):',
                                                               '<b>Servings per Package (EACH):</b>',
                                                               '<b>Servings per Package (INNER):</b>']:
                    if final_content_list[p][1] in ['Legal_Designation', 'LOCATION_OF_ORIGIN', 'VARIANT',
                                                    'Servings Per Package', 'Serving Size',
                                                    '<b>Serving size/ حجم الحصة</b>', 'BRAND_NAME',
                                                    'Serving Per Container', 'NET_CONTENT_STATEMENT']:
                        split_cnt = final_content_list[p][0][d].split('\n')
                        for texts in split_cnt:
                            if texts.replace('<b> </b>', '').strip() != '':
                                lang = classify(texts)[0]
                                if final_content_list[p][1] in gen_cate_dic:
                                    gen_cate_dic[final_content_list[p][1]].append(
                                        {lang: texts.replace('<b> </b>', '').strip()})
                                else:
                                    gen_cate_dic[final_content_list[p][1]] = [
                                        {lang: texts.replace('<b> </b>', '').strip()}]
                    else:
                        if final_content_list[p][0][d].strip() != '':
                            lang = classify(final_content_list[p][0][d])[0]
                            if final_content_list[p][1] in gen_cate_dic:
                                gen_cate_dic[final_content_list[p][1]].append(
                                    {lang: final_content_list[p][0][d].replace('<b> </b>', '')})
                            else:
                                gen_cate_dic[final_content_list[p][1]] = [
                                    {lang: final_content_list[p][0][d].replace('<b> </b>', '')}]

        elif len(final_content_list[p][0]) == 1:
            lang = classify(final_content_list[p][0][0])[0]
            if final_content_list[p][0][0].strip() != '':
                if final_content_list[p][1] in gen_cate_dic:
                    gen_cate_dic[final_content_list[p][1]].append(
                        {lang: final_content_list[p][0][0].replace('<b> </b>', '')})
                else:
                    gen_cate_dic[final_content_list[p][1]] = [
                        {lang: final_content_list[p][0][0].replace('<b> </b>', '')}]

    for i in range(0, len(table_content_list_all)):
        if 'brand' in table_content_list_all[i][0].lower():
            brand_indx = i
            if len(table_content_list_all[brand_indx]) == 2:
                if table_content_list_all[brand_indx + 1][0].strip() == '':
                    if table_content_list_all[brand_indx + 1][1].strip() != '':
                        lang = classify(table_content_list_all[brand_indx + 1][1])[0]
                        if "BRAND_NAME" in gen_cate_dic:
                            gen_cate_dic["BRAND_NAME"].append({lang: table_content_list_all[brand_indx + 1][1]})
                        else:
                            gen_cate_dic["BRAND_NAME"] = [{lang: table_content_list_all[brand_indx + 1][1]}]

    '''  # Allergen statement not needed as we grouping it under ingredient declaration
    for k3, v3 in ingre_new_dic.items():
        if k3 == 'ALLERGEN_STATEMENT':
            for cnt_list in v3:
                for lang, texts in cnt_list.items():
                    if k3 in gen_cate_dic:
                        gen_cate_dic[k3].append({lang: texts})
                    else:
                        gen_cate_dic[k3] = [{lang: texts}]
    '''

    ingre_new_dic.pop('ALLERGEN_STATEMENT', None)
    t2 = time.time()
    print(f'Complted in {t2 - t1} secs')
    return table_content_list, gen_cate_dic, nutri_para_list, ingre_dic, thai_nutrition, indo_nutrition, en_nutrition


# In[6]:


def multi_nutri_text(table_content_list):
    en_nutrition = []
    for k in range(0, len(table_content_list)):
        split_cnt = re.split('/|,', table_content_list[k][0])  # /|,
        #         if "shelf life" not in split_cnt[0].lower():
        classified_output = classifier.predict(laser.embed_sentences(
            split_cnt[0].replace('&lt;', '').replace('&gt;', '').replace('</b>', '').replace('<b>', ''), lang='en'))
        probability1 = classifier.predict_proba(laser.embed_sentences(
            split_cnt[0].replace('&lt;', '').replace('&gt;', '').replace('</b>', '').replace('<b>', ''), lang='en'))
        probability1.sort()
        prob1 = probability1[0][-1]
        if prob1 > 0.75:
            classified_output = classified_output[0]
        else:
            classified_output = 'UNMAPPED'

        if classified_output == 'EN_NUTRITION':
            en_nutrition.append(table_content_list[k])
    return en_nutrition


# In[7]:


def paragraph(file):
    document = Document(file)

    paragraph_txt = []
    all_paras = document.paragraphs
    for para in all_paras:
        paragraph_txt.append(para.text.strip().replace('<', '&lt;').replace('>', '&gt;'))

    paragraph_txt = [txt.strip() for txt in paragraph_txt if txt]

    para_dic = {}
    for txt in paragraph_txt:
        #         if txt[1] in "*GDA":
        #             value_gda = txt[1:]
        #             para_dic.setdefault("NUTRITION_TABLE_CONTENT", []).append({classify(value_gda)[0]:value_gda.strip()})
        classified_output = classifier.predict(
            laser.embed_sentences(txt.replace('</b>', '').replace('<b>', ''), lang='en'))
        probability1 = classifier.predict_proba(
            laser.embed_sentences(txt.replace('</b>', '').replace('<b>', ''), lang='en'))
        probability1.sort()
        prob1 = probability1[0][-1]
        if prob1 > 0.75:
            classified_output = classified_output[0]
        else:
            classified_output = 'UNMAPPED'

        if classified_output == "NUTRITION_TABLE_CONTENT":

            lang = classify(txt)[0]
            if "NUTRITION_TABLE_CONTENT" in para_dic:
                para_dic["NUTRITION_TABLE_CONTENT"].append({lang: txt})
            else:
                para_dic["NUTRITION_TABLE_CONTENT"] = [{lang: txt}]

    return para_dic


# In[8]:


def remove_duplicates(dictionary):
    final_cleaned_dict = {}
    for category, value_list in dictionary.items():
        final_cleaned_dict[category] = sorted(
            list({frozenset(list_element.keys()): list_element for list_element in value_list}.values()),
            key=lambda d: list(d.keys()))
    return final_cleaned_dict


# In[9]:


def nutrition(dictionary):
    nutri_dic = {}
    for keys, value in dictionary.items():
        split_cnt = re.split(r'/|,', keys)
        classified_output_1 = classifier_key.predict(laser.embed_sentences(
            split_cnt[0].replace('\xa0', '').replace('&lt;', '').replace('&gt;', '').replace('</b>', '').replace('<b>',
                                                                                                                 '').replace(
                '\x95', ''), lang='en'))
        probability_1 = classifier_key.predict_proba(laser.embed_sentences(
            split_cnt[0].replace('\xa0', '').replace('&lt;', '').replace('&gt;', '').replace('\x95', '').replace('</b>',
                                                                                                                 '').replace(
                '<b>',
                ''),
            lang='en'))
        probability_1.sort()
        prob_1 = probability_1[0][-1]
        if prob_1 > 0.65:
            classified_output_key = classified_output_1[0]
        else:
            classified_output_key = 'UNMAPPED'

        if classified_output_key != 'None':

            if classified_output_key in nutri_dic:
                if isinstance(value, list):
                    nutri_dic[classified_output_key].extend(value)
                elif isinstance(value, dict):
                    nutri_dic[classified_output_key].append(value)
            else:
                nutri_dic[classified_output_key] = value

        else:
            if 'calcium' in keys.lower():
                if 'Calcium' in nutri_dic:
                    if isinstance(value, list):
                        nutri_dic['Calcium'].extend(value)
                    elif isinstance(value, dict):
                        nutri_dic['Calcium'].append(value)
                else:
                    nutri_dic['Calcium'] = value
            else:
                if keys in nutri_dic:
                    if isinstance(value, list):
                        nutri_dic[keys].extend(value)
                    elif isinstance(value, dict):
                        nutri_dic[keys].append(value)
                else:
                    nutri_dic[keys] = value

    return nutri_dic


# In[10]:


def serving_dictionary(dic, gen_cate_dic):
    for k, v in dic.items():
        classified_output = classifier_key.predict(
            laser.embed_sentences(k.replace('\xa0', '').replace('</b>', '').replace('<b>', ''), lang='en'))
        probability1 = classifier_key.predict_proba(
            laser.embed_sentences(k.replace('\xa0', '').replace('</b>', '').replace('<b>', ''), lang='en'))
        probability1.sort()
        prob1 = probability1[0][-1]
        if prob1 > 0.75:
            classified_output = classified_output[0]
        else:
            classified_output = 'UNMAPPED'

        if classified_output != 'None':
            if classified_output in gen_cate_dic:

                gen_cate_dic[classified_output].append(v)
            else:
                gen_cate_dic[classified_output] = v
        else:
            if k in gen_cate_dic:

                gen_cate_dic[k].append(v)
            else:
                gen_cate_dic[k] = v

    return gen_cate_dic


# In[11]:


def nutrition_para_format(nutri_para_list, gen_cate_dic, table_content_list):
    content_list = []
    for i in nutri_para_list:
        list_of_cnt = i.replace('</b>', '').replace('<b>', '').replace('<', '&lt;').replace('>', '&gt;').split(',')
        content_list.append(list_of_cnt)

    nutri_dic = []
    for k1 in range(0, len(content_list)):

        nutrit_para_dic = {}
        #         nutri_dic_new={}
        d_regex_list = []

        for i in content_list[k1]:
            rege = ('').join(re.findall(r'[<.a-zA-Z0-9-(&lt;)(%)]', str(i)))
            d_regex_list.append(rege)

        d_regex_new_list = []
        for j in d_regex_list:
            reg1 = re.findall(r'((&lt;)?\s?(\d?\d)(\.\d+)?\s?(mg|kj|g|%|mcg))', str(j))
            d_regex_new_list.append(reg1)
            #             print(d_regex_new_list)

        #         print(d_regex_new_list)

        for l5 in range(0, len(d_regex_new_list)):
            for m5 in range(0, len(d_regex_new_list[l5])):
                for keys in nutri_keys_para:
                    if keys in content_list[k1][l5].lower().replace('\xa0', '').replace('\n', ''):
                        #                         print(keys)
                        if '%' in d_regex_new_list[l5][m5][0]:
                            if keys in nutrit_para_dic:
                                nutrit_para_dic[keys].append({'PDV': {'en': d_regex_new_list[l5][m5][0]}})
                            else:
                                nutrit_para_dic[keys] = [{'PDV': {'en': d_regex_new_list[l5][m5][0]}}]
                        else:
                            if keys in nutrit_para_dic:
                                nutrit_para_dic[keys].append({'Value': {'en': d_regex_new_list[l5][m5][0]}})
                            else:
                                nutrit_para_dic[keys] = [{'Value': {'en': d_regex_new_list[l5][m5][0]}}]
        #         print(nutrit_para_dic)
        nutri_dic_new = nutrition(nutrit_para_dic)

        nutri_dic.append(nutri_dic_new)

    serving_list = []
    serving_key = ["servings per container", "amount per serving"]
    for i in range(0, len(table_content_list)):
        for k in serving_key:
            if k in table_content_list[i][0].lower():
                serving_list.append(table_content_list[i][0].replace('\xa0', ''))

    final_serving_list = []
    for j in range(0, len(serving_list)):
        split_list = serving_list[j].split(',')
        final_serving_list.append(split_list)

    final_serving_list = sum(final_serving_list, [])

    para_regex_list = []
    para_regex_new_list = []
    for i in final_serving_list:
        #     reg = re.findall(r'(<?\s+(\s?\d)(\.\d+)?\s?(mg|kj|g|%|mcg))|(<?\s+(\.\d+)?\s?(mg|kj|g|%|mcg))',str(i))
        para_reg = ('').join(re.findall(r'[<.a-zA-Z0-9-(%)]', str(i)))
        para_regex_list.append(para_reg)

        # regex_list

    for j in para_regex_list:
        para_reg1 = re.findall(r'(<?\s?(\d?\d)(\.\d+)?\s?(mg|kj|g|%|mcg))|(\d?\d?\d)', str(j))
        para_regex_new_list.append(para_reg1)

    para_tuple_list = []
    for m in para_regex_new_list:
        para_tuple_list.append([(tuple(int(x) if x.isdigit() else x for x in _ if x)) for _ in m])

    serving_dic_items = {}
    serving_dic = {}
    for l in range(0, len(para_tuple_list)):
        for m in range(0, len(para_tuple_list[l])):
            for k in nutri_keys_para:
                if k in final_serving_list[l].replace('\x95', '').replace('\xa0', '').lower():
                    if '%' in str(para_tuple_list[l][m][0]):
                        if k in serving_dic_items:
                            serving_dic_items[k].append({'en': str(para_tuple_list[l][m][0])})
                        else:
                            serving_dic_items[k] = [{'en': str(para_tuple_list[l][m][0])}]
                    else:
                        if k in serving_dic_items:
                            serving_dic_items[k].append({'en': str(para_tuple_list[l][m][0])})
                        else:
                            serving_dic_items[k] = [{'en': str(para_tuple_list[l][m][0])}]

    gen_cate_dic = serving_dictionary(serving_dic_items, gen_cate_dic)
    gen_cate_dic.pop('Calories', None)
    return gen_cate_dic, nutri_dic


# In[12]:


def non_table_nutrition(table_content_list):
    #     print("new_non")
    new_non_table = []
    for i in range(0, len(table_content_list)):
        for j in range(0, len(table_content_list[i])):
            if 'nutrition information' in table_content_list[i][j].lower() or 'nutritional information' in \
                    table_content_list[i][j].lower():
                if not new_non_table:
                    if 'serving size' in table_content_list[i + 1][0].lower():
                        new_non_table.append(table_content_list[i + 2])
                    else:
                        new_non_table.append(table_content_list[i + 1])

    keys_list = []
    nutri_values = []
    for j in range(0, len(new_non_table)):
        for k in range(0, len(new_non_table[j])):
            #             print(new_non_table[j][k])
            if 'energy' in new_non_table[j][k].lower():
                keys_list.append(
                    new_non_table[j][k].replace('</b>', '').replace('<b>', '').replace('\xa0', '').split('\n'))
            else:
                if '\n' in new_non_table[j][k].lower():
                    nutri_values.append(new_non_table[j][k].replace('</b>', '').replace('<b>', '').split('\n'))
    nutri_keys = []
    for i in range(0, len(keys_list)):
        for j in range(0, len(keys_list[i])):
            split = keys_list[i][j].split('/')[0]
            nutri_keys.append(split)

    new_nutri_dic = {}
    for j1 in range(0, len(nutri_values)):
        # if len(nutri_keys[0]) == len(nutri_values[j1]):
        if len(nutri_keys) == len(nutri_values[j1]):
            for j2 in range(0, len(nutri_values[j1])):
                if nutri_values[j1][j2].strip() != '':
                    # if nutri_keys[0][j2] in new_nutri_dic:
                    if nutri_keys[j2] in new_nutri_dic:
                        if '%' in nutri_values[j1][j2].strip():
                            # new_nutri_dic[nutri_keys[0][j2]].append({'PDV':{'en':nutri_values[j1][j2].strip()}})
                            new_nutri_dic[nutri_keys[j2]].append({'PDV': {'en': nutri_values[j1][j2].strip()}})
                        else:
                            if 'kcal' in nutri_values[j1][j2].strip().lower() and 'energy' not in nutri_keys[
                                j2].lower():
                                new_nutri_dic['calories'].append({'Values': {'en': nutri_values[j1][j2].strip()}})
                            else:
                                # new_nutri_dic[nutri_keys[0][j2]].append({'Values':{'en':nutri_values[j1][j2].strip()}})
                                new_nutri_dic[nutri_keys[j2]].append({'Values': {'en': nutri_values[j1][j2].strip()}})
                    else:
                        if '%' in nutri_values[j1][j2].strip():
                            new_nutri_dic[nutri_keys[j2]] = [{'PDV': {'en': nutri_values[j1][j2].strip()}}]
                        else:

                            if 'kcal' in nutri_values[j1][j2].strip().lower() and 'energy' not in nutri_keys[
                                j2].lower():
                                if 'calories' in new_nutri_dic:
                                    new_nutri_dic['calories'].append({'Values': {'en': nutri_values[j1][j2].strip()}})
                                else:
                                    new_nutri_dic['calories'] = [{'Values': {'en': nutri_values[j1][j2].strip()}}]
                            else:
                                # new_nutri_dic[nutri_keys[0][j2]]= [{'Values':{'en':nutri_values[j1][j2].strip()}}]
                                new_nutri_dic[nutri_keys[j2]] = [{'Values': {'en': nutri_values[j1][j2].strip()}}]

    nutri_dic = nutrition(new_nutri_dic)

    return nutri_dic


# In[13]:


def arabic_table_nutrition(table_content_list, gen_cate_dic):
    start_indx = 0
    end_indx = 0
    for i in range(0, len(table_content_list)):
        if 'nutrition facts' in table_content_list[i][0].lower():
            start_indx = i + 1
        elif "% daily value" in table_content_list[i][0].lower():
            end_indx = i
            break
    new_list = []
    for j in range(start_indx, end_indx):
        new_list.append(table_content_list[j])

    dupli_list = []
    for i in new_list:
        dupli_list.append(('').join(i).replace('</b>', '').replace('<b>', '').replace('<', '&lt;').replace('>', '&gt;'))

    regex_list_ = []
    for i in dupli_list:
        #     reg = re.findall(r'(<?\s+(\s?\d)(\.\d+)?\s?(mg|kj|g|%|mcg))|(<?\s+(\.\d+)?\s?(mg|kj|g|%|mcg))',str(i))
        reg_ = ('').join(re.findall(r'[<.a-zA-Z0-9-(&lt;)(%)]', str(i).replace(',', '.')))
        regex_list_.append(reg_)
        # regex_list

    regex_new_list_ = []
    for j in regex_list_:
        #     reg1_ = re.findall(r'(<?\s?(\d?\d?\d)(\.\d+)?\s?(mg|kj|g|%|mcg))|(\d?\d?\d)',str(j).replace(',','.'))
        reg1_ = re.findall(r'((&lt;)?\s?(\d+)(\.\d+)?(\,\d+)?\.?\s?(mg|kj|g|%|mcg))|(\d?\d?\d)',
                           str(j).replace(',', '.'))
        regex_new_list_.append(reg1_)

    tuple_list = []
    for m in regex_new_list_:
        tuple_list.append([(tuple(int(x) if x.isdigit() else x for x in _ if x)) for _ in m])

    regex_new_list_ = tuple_list

    nutri_arabic_dic = {}
    serving_dic = {}
    for l in range(0, len(regex_new_list_)):
        for m in range(0, len(regex_new_list_[l])):
            for k in nutri_keys_para:
                if k in dupli_list[l].lower().replace('\x95', '').replace('\xa0', ''):
                    if '%' in str(regex_new_list_[l][m][0]):
                        if k in nutri_arabic_dic:
                            nutri_arabic_dic[k].append({'PDV': {'en': str(regex_new_list_[l][m][0])}})
                        else:
                            nutri_arabic_dic[k] = [{'PDV': {'en': str(regex_new_list_[l][m][0])}}]
                    else:
                        if k in nutri_arabic_dic:
                            nutri_arabic_dic[k].append({'Value': {'en': str(regex_new_list_[l][m][0])}})
                        else:
                            nutri_arabic_dic[k] = [{'Value': {'en': str(regex_new_list_[l][m][0])}}]

    duplicate_dic = remove_duplicates(nutri_arabic_dic)

    nutri_dic = nutrition(duplicate_dic)

    serving_dic_items = {}
    serving_dic = {}
    serv_list = ['serving size', 'servings per container', 'servings per package']
    #         for k6,v6 in duplicate_dic.items():
    #             for ke in serv_list:
    #                 if ke in k6.lower():
    #                     if ke in serving_dic_items:
    #                           serving_dic_items[ke].append(v6)
    #                     else:
    #                         serving_dic_items[ke] = v6
    for k, v in duplicate_dic.items():
        for txt in v:
            for k7, v7 in txt.items():
                for ke in serv_list:
                    if ke in k.lower():
                        if ke in serving_dic_items:
                            serving_dic_items[ke].append(v7)
                        else:
                            serving_dic_items[ke] = v7

    gen_cate_dic = serving_dictionary(serving_dic_items, gen_cate_dic)

    nutri_dic.pop('Serving Size', None)
    nutri_dic.pop('Serving Per Container', None)
    nutri_dic.pop('Servings Per Package', None)
    nutri_dic.pop('Servings Per Container', None)

    return nutri_dic, gen_cate_dic


# In[14]:


def thai_nutri_dic(thai_nutrition):
    thai_nutrition_list = [('').join(i).replace('\xa0', '').replace('\n', '') for i in thai_nutrition]

    reg_list = []
    for i in thai_nutrition_list:
        reg = re.findall(r'((&lt;)?\s?(\d+)(\.\d+)?\s?(mg|kj|g|ก|มก|%|mcg))',
                         i.replace('\n', '').replace('</b>', '').replace('<b>', ''))
        reg_list.append(reg)

    thai_nutrition_key = []
    for a in range(0, len(thai_nutrition)):
        key_reg = ('').join(re.findall(r'([^(&lt;)?\s?(\d+)(\.\d+)?\s?(mg|kj|g|%|mcg)])',
                                       thai_nutrition[a][0].replace('มก', '').replace('ก', '').replace('</b>',
                                                                                                       '').replace(
                                           '<b>', '')))
        thai_nutrition_key.append(key_reg)

    thai_nutri_dic = {}
    for j in range(0, len(reg_list)):
        for n in range(0, len(reg_list[j])):
            if reg_list[j][n][0] != '':
                if thai_nutrition_key[j] in thai_nutri_dic:
                    if '%' not in reg_list[j][n][0]:
                        thai_nutri_dic[thai_nutrition_key[j]].append({'Value': {'en': reg_list[j][n][0].strip()}})
                    else:
                        thai_nutri_dic[thai_nutrition_key[j]].append({'PDV': {'en': reg_list[j][n][0].strip()}})
                else:
                    if '%' not in reg_list[j][n][0]:
                        thai_nutri_dic[thai_nutrition_key[j]] = [{'Value': {'en': reg_list[j][n][0].strip()}}]
                    else:
                        thai_nutri_dic[thai_nutrition_key[j]] = [{'PDV': {'en': reg_list[j][n][0].strip()}}]

    return thai_nutri_dic


# In[15]:


def indo_en_nutrition(indo_nutrition):
    indo_nutri_dic = {}
    for l in range(0, len(indo_nutrition)):
        if len(indo_nutrition[l]) > 1:
            for m in range(1, len(indo_nutrition[l])):
                if indo_nutrition[l][m].replace('</b>', '').replace('<b>', '').strip() not in ['', '(kcal)']:
                    if indo_nutrition[l][0].replace('\xa0', '').strip() in indo_nutri_dic:
                        if '%' not in indo_nutrition[l][m]:
                            indo_nutri_dic[indo_nutrition[l][0].replace('\xa0', '').strip()].append({'Value': {
                                'en': indo_nutrition[l][m].replace('\n', '').replace('\xa0', '').replace('</b>',
                                                                                                         '').replace(
                                    '<b>', '').replace('<', '&lt;').replace('>', '&gt;').strip()}})
                        else:
                            indo_nutri_dic[indo_nutrition[l][0].replace('\xa0', '').strip()].append({'PDV': {
                                'en': indo_nutrition[l][m].replace('\n', '').replace('\xa0', '').replace('</b>',
                                                                                                         '').replace(
                                    '<b>', '').replace('<', '&lt;').replace('>', '&gt;').strip()}})
                    else:
                        if '%' not in indo_nutrition[l][m]:
                            indo_nutri_dic[indo_nutrition[l][0].replace('\xa0', '').strip()] = [{'Value': {
                                'en': indo_nutrition[l][m].replace('\n', '').replace('\xa0', '').replace('</b>',
                                                                                                         '').replace(
                                    '<b>', '').replace('<', '&lt;').replace('>', '&gt;').strip()}}]
                        else:
                            indo_nutri_dic[indo_nutrition[l][0].replace('\xa0', '').strip()] = [{'PDV': {
                                'en': indo_nutrition[l][m].replace('\n', '').replace('\xa0', '').replace('</b>',
                                                                                                         '').replace(
                                    '<b>', '').replace('<', '&lt;').replace('>', '&gt;').strip()}}]

    return indo_nutri_dic


# In[16]:


def remove_specific_duplicates(dictionary, spec_cate):
    final_spec_dict = {}
    for category, value_list in dictionary.items():
        if category in spec_cate:
            final_spec_dict[category] = sorted(
                list({frozenset(list_element.values()): list_element for list_element in value_list}.values()),
                key=lambda d: list(d.keys()))

        else:
            if category in final_spec_dict:
                final_spec_dict[category].append(value_list)
            else:
                final_spec_dict[category] = value_list

    return final_spec_dict

def get_input(input_file, input_docx_location):
    if input_file.startswith('\\'):
        print('connecting to SMB share')
        try:
            with smbclient.open_file(r"{}".format(input_file), mode='rb', username=smb_username,
                                     password=smb_password) as f:
                with open(input_docx_location, 'wb') as pdf:
                    pdf.write(f.read())
                print('file found')
        except:
            smbclient.reset_connection_cache()
            with smbclient.open_file(r"{}".format(input_file), mode='rb', username=smb_username,
                                     password=smb_password) as f:
                with open(input_docx_location, 'wb') as pdf:
                    pdf.write(f.read())
                print('file found')
        finally:
            smbclient.reset_connection_cache()
        return input_docx_location
    else:
        return document_location + input_file

def mondelez_mea_word_main(file):
    t5 = time.time()
    nutrition_dic = []
    temp_directory = tempfile.TemporaryDirectory(dir=document_location)
    input_docx_location = f'{temp_directory.name}/input_docx.docx'
    file = get_input(file, input_docx_location)
    flag = 0
    ## bold_list = bold_text_list(file)
    table_content_list, gen_cate_dic, nutri_para_list, ingre_dic, thai_nutrition, indo_nutrition, en_nutrition = preprocess(
        file)
    grt = paragraph(file)
    ## para_dic = paragraph(file)
    if nutri_para_list:
        gen_cate_dic, nutrition_dic = nutrition_para_format(nutri_para_list, gen_cate_dic, table_content_list)


    elif not en_nutrition and not thai_nutrition and not indo_nutrition:
        for i in range(0, len(table_content_list)):
            for j in range(0, len(table_content_list[i])):
                if 'nutrition information' in table_content_list[i][j].lower() or 'nutritional information' in \
                        table_content_list[i][j].lower():
                    flag = 1
                elif 'nutrition facts' in table_content_list[i][j].lower():
                    flag = 2
                else:
                    pass

        if flag == 1:
            nutrition_dic = non_table_nutrition(table_content_list)

        elif flag == 2:
            nutrition_dic, gen_cate_dic = arabic_table_nutrition(table_content_list, gen_cate_dic)

        else:
            pass

    #         if not non_table:
    #             nutrition_dic, gen_cate_dic = arabic_table_nutrition(table_content_list,gen_cate_dic)
    #         else:

    #             nutrition_dic = non_table_nutrition(non_table)

    else:
        # nutrition_dic = []
        if en_nutrition:
            if len(en_nutrition) > 3:
                nutri_dic = indo_en_nutrition(en_nutrition)
                final_dict = nutrition(nutri_dic)
                nutrition_dic.append(final_dict)
        if thai_nutrition:
            if len(thai_nutrition) > 3:
                nutri_dic = thai_nutri_dic(thai_nutrition)
                final_dict = nutrition(nutri_dic)
                nutrition_dic.append(final_dict)
        if indo_nutrition:
            for s1 in range(0, len(table_content_list)):
                for s2 in range(0, len(table_content_list[s1])):
                    if 'nutrition facts' in table_content_list[s1][s2].lower():
                        flag = 4
            if flag == 4:
                nutrition_dic, gen_cate_dic = arabic_table_nutrition(table_content_list, gen_cate_dic)
            else:
                if len(indo_nutrition) > 3:
                    nutri_dic = indo_en_nutrition(indo_nutrition)
                    final_dict = nutrition(nutri_dic)
                    nutrition_dic.append(final_dict)

    gen_cate_dic = remove_specific_duplicates(gen_cate_dic, ['NUTRITION_TABLE_CONTENT'])
    over_all_dic = {**gen_cate_dic, **ingre_dic, **grt}

    ## Newly added script for multilanguage nutrition in same text.
    if not nutrition_dic:
        nutrition_dic = []
        en_nutrition = multi_nutri_text(table_content_list)
        if len(en_nutrition) > 3:
            nutri_dic = indo_en_nutrition(en_nutrition)
            final_dict = nutrition(nutri_dic)
            nutrition_dic.append(final_dict)

    final_dic = {}
    if 'NUTRITION_FACTS' in final_dic:
        if isinstance(nutrition_dic, list):
            final_dic['NUTRITION_FACTS'].extend(nutrition_dic)
        else:
            final_dic['NUTRITION_FACTS'].append(nutrition_dic)
    else:
        if isinstance(nutrition_dic, list):
            final_dic['NUTRITION_FACTS'] = nutrition_dic
        else:
            final_dic['NUTRITION_FACTS'] = [nutrition_dic]
    #     final_dic['NUTRITION_FACTS'] = nutrition_dic
    t6 = time.time()
    print(f'Finished in {t6 - t5}seconds')
    return {**final_dic, **over_all_dic}


