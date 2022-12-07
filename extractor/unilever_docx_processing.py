from docx import Document
from langid import classify
from sklearn.neural_network import MLPClassifier
import time
from laserembeddings import Laser
import warnings
import smbclient
import tempfile

from .utils import tag_convert
from environment import MODE

if MODE == 'local':
    from .local_constants import *
else:
    from .dev_constants import *

warnings.filterwarnings("ignore")
import joblib

laser = Laser(path_to_bpe_codes, path_to_bpe_vocab, path_to_encoder)

# model_location = r"/Users/sakthivel/Documents/SGK/Unilever/Dataset/unilever_docx_model.sav"
classifier = joblib.load(unilever_docx_model_location)

def get_input(input_file,input_docx_location):
    if input_file.startswith('\\'):
        print('connecting to SMB share')
        try:
            with smbclient.open_file(r"{}".format(input_file), mode='rb', username=smb_username,
                                     password=smb_password) as f:
                with open(input_docx_location, 'wb') as pdf:
                    pdf.write(f.read())
                print('file found')
        except:
            smbclient.reset_connection_cache()
            with smbclient.open_file(r"{}".format(input_file), mode='rb', username=smb_username,
                                     password=smb_password) as f:
                with open(input_docx_location, 'wb') as pdf:
                    pdf.write(f.read())
                print('file found')
        finally:
            smbclient.reset_connection_cache()
        return input_docx_location
    else:
        return document_location + input_file

def table_content(file):
    document = Document(file)
    table_content_list = []
    for table in document.tables:
        for row in table.rows:
            row_list = []
            for cell in row.cells:
                row_list.append(str(cell.text))
            table_content_list.append(row_list)
    return table_content_list

def table_content_dic(table_content_list):
    final_dic = {}
    for value in table_content_list:
        lang = classify(value[0])[0]
        value_converted = tag_convert(value[0])
        if "INTERNAL_PACKAGE_IDENTIFIER" in final_dic:
            final_dic["INTERNAL_PACKAGE_IDENTIFIER"].append({lang: value_converted})
        else:
            final_dic["INTERNAL_PACKAGE_IDENTIFIER"] = [{lang: value_converted}]
    return final_dic

def paragrapgh_txt(file):
    document = Document(file)
    paragraph_txt = []
    for para in document.paragraphs:
        paragraph_txt.append(para.text.strip())
    return paragraph_txt


def classifier_dic(paragraph_txt_list):
    dic = {}
    for items in paragraph_txt_list:
        if items.strip() != "":
            classified_output = classifier.predict(
                laser.embed_sentences(items.replace('<b>', '').replace('</b>', ''), lang='en'))
            probability1 = classifier.predict_proba(
                laser.embed_sentences(items.replace('<b>', '').replace('</b>', ''), lang='en'))
            probability1.sort()
            prob1 = probability1[0][-1]
            if prob1 > 0.75:
                classified_output = classified_output[0]
            else:
                classified_output = 'UNMAPPED'
            lang = classify(items)[0]
            items = tag_convert(items)                              # < and > symbol replace

            if classified_output in dic:
                dic[classified_output].append({lang: items})
            else:
                dic[classified_output] = [{lang: items}]
    #         print(prob1,classified_output)
    return dic


def unilever_docx_main(input_file):
    temp_directory = tempfile.TemporaryDirectory(dir=document_location)
    input_docx_location = f'{temp_directory.name}/input_docx.docx'
    file = get_input(input_file, input_docx_location)
    t1 = time.time()
    if 'pm code' in input_file.lower():
        table_content_list = table_content(file)
        table_content_dictionary = table_content_dic(table_content_list)
        t2 = time.time()
        print(f'Finished in {t2 - t1}seconds')
        return table_content_dictionary

    elif 'summary' in input_file.lower() or 'hotline' in input_file.lower():
        print('inside summary document')
        paragraph_list = paragrapgh_txt(file)
        final_dic = classifier_dic(paragraph_list)
        t2 = time.time()
        print(f'Finished in {t2 - t1}seconds')
        return final_dic
    else:
        pass